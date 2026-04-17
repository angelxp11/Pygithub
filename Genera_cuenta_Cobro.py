import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from tkcalendar import DateEntry
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx import Document
from num2words import num2words
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert
import threading
import os
import time
import subprocess
import sys
import pythoncom
import io

# ==============================
# MANEJO DE STDOUT/STDERR (PyInstaller --noconsole)
# ==============================
if sys.stdout is None:
    sys.stdout = io.StringIO()
if sys.stderr is None:
    sys.stderr = io.StringIO()

# ==============================
# FUNCIONES
# ==============================
def recurso_path(rel_path):
    """Retorna la ruta correcta para recursos, compatible con PyInstaller"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, rel_path)

def set_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)

    tcPr.append(tcBorders)

def formato_cop(valor):
    return "${:,.0f}".format(valor).replace(",", ".")

def leer_txt(contenido):
    lineas = [l.strip() for l in contenido.split("\n") if l.strip()]

    if len(lineas) % 3 != 0:
        raise Exception("Formato incorrecto del TXT")

    fechas, actividades, horas = [], [], []

    for i in range(0, len(lineas), 3):
        fechas.append(lineas[i])
        actividades.append(lineas[i+1])
        horas.append(lineas[i+2])

    return fechas, actividades, horas



# ==============================
# DRAG & DROP
# ==============================
def drop(event):
    try:
        rutas = root.tk.splitlist(event.data)
        ruta = rutas[0]

        if not ruta.lower().endswith(".txt"):
            messagebox.showerror("Error", "Solo archivos .txt")
            return

        with open(ruta, "r", encoding="utf-8") as f:
            contenido = f.read()

        txt_input.delete("1.0", tk.END)
        txt_input.insert(tk.END, contenido)

        estado.set(f"TXT cargado ✔\n{ruta}")

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo cargar:\n{e}")
    try:
        rutas = root.tk.splitlist(event.data)  # 🔥 Maneja múltiples archivos correctamente
        ruta = rutas[0]

        if not ruta.lower().endswith(".txt"):
            messagebox.showerror("Error", "Solo se permiten archivos TXT")
            return

        with open(ruta, "r", encoding="utf-8") as f:
            contenido = f.read()

        txt_input.delete("1.0", tk.END)
        txt_input.insert(tk.END, contenido)

        estado.set("TXT cargado por arrastre ✔")

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo cargar:\n{e}")
    ruta = event.data.strip("{}")  # Windows fix
    try:
        with open(ruta, "r", encoding="utf-8") as f:
            txt_input.delete("1.0", tk.END)
            txt_input.insert(tk.END, f.read())
        estado.set("TXT cargado por arrastre ✔")
    except:
        messagebox.showerror("Error", "Archivo no válido")

# ==============================
# GENERAR
# ==============================
def generar():
    # 🔒 Bloquear botón
    btn_generar.config(state="disabled")
    
    # Inicializar COM para threads
    pythoncom.CoInitialize()
    
    try:
        inicio = time.time()  # ⏱️ iniciar contador

        progreso.set(10)
        estado.set("Leyendo datos...")

        contenido = txt_input.get("1.0", tk.END)
        valor_total = int(valor_entry.get())

        fecha_dt = calendario.get_date()
        fecha_texto = fecha_dt.strftime("%d de %B %Y")

        fechas, actividades, horas = leer_txt(contenido)

        progreso.set(30)
        estado.set("Generando DOCX...")

        ruta_docx = recurso_path("Formatocuentadecobro.docx")
        doc = Document(ruta_docx)

        valor_letras = num2words(valor_total, lang='es').upper() + " PESOS"

        for p in doc.paragraphs:
            if "(VALOR TOTAL NUMERO)" in p.text:
                p.text = p.text.replace("(VALOR TOTAL NUMERO)", formato_cop(valor_total))
            if "(VALOR TOTAL LETRAS EN PESOS)" in p.text:
                p.text = p.text.replace("(VALOR TOTAL LETRAS EN PESOS)", valor_letras)
            if "(FECHA SUPERIOR)" in p.text:
                p.text = p.text.replace("(FECHA SUPERIOR)", fecha_texto)

        tabla = doc.tables[0]

        # encabezado
        for cell in tabla.rows[0].cells:
            set_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # ==============================
        # FILAS
        # ==============================
        filas_valor = []

        for i in range(len(fechas)):
            row = tabla.add_row().cells

            row[0].text = fechas[i].strip()
            row[1].text = actividades[i].strip()
            row[2].text = f"{horas[i].strip()} Horas"
            row[3].text = ""

            filas_valor.append(row[3])

            for j, cell in enumerate(row):
                set_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.size = Pt(9)

        # ==============================
        # 🔥 COMBINAR COLUMNA VALOR TOTAL
        # ==============================
        if filas_valor:
            celda_principal = filas_valor[0]

            for celda in filas_valor[1:]:
                celda_principal = celda_principal.merge(celda)

            celda_principal.text = formato_cop(valor_total)

            # centrado horizontal + vertical
            celda_principal.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for p in celda_principal.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(12)

            set_borders(celda_principal)

        # ==============================
        # AJUSTAR TAMAÑO PARA QUE QUEPA EN UNA HOJA
        # ==============================
        from docx.oxml import parse_xml
        
        tbl = tabla._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Ajustar ancho de tabla al 100% de la página
        tbl_width = parse_xml(r'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="5000" w:type="auto"/>')
        
        # Remover anchos anteriores si existen
        for width_elem in tblPr.findall(qn('w:tblW')):
            tblPr.remove(width_elem)
        
        tblPr.append(tbl_width)

        # ==============================
        # NOMBRE AUTOMÁTICO
        # ==============================
        fecha_archivo = fecha_dt.strftime("%Y-%m-%d")

        # 📂 ruta donde se guardará
        carpeta_salida = os.getcwd()

        docx_file = os.path.join(carpeta_salida, f"Cuenta_{fecha_archivo}.docx")
        pdf_file = os.path.join(carpeta_salida, f"Cuenta_{fecha_archivo}.pdf")

        if os.path.exists(docx_file):
            os.remove(docx_file)

        doc.save(docx_file)

        progreso.set(70)
        estado.set("DOCX listo ✔ - Generando PDF...")

        convert(docx_file, pdf_file)

        progreso.set(100)

        fin = time.time()  # ⏱️ fin
        tiempo = round(fin - inicio, 2)

        estado.set(f"PDF generado ✔ en {tiempo} segundos")

        # 📂 ABRIR CARPETA AUTOMÁTICAMENTE (Windows)
        subprocess.Popen(f'explorer "{carpeta_salida}"')

        messagebox.showinfo(
            "Éxito",
            f"Archivos generados:\n{docx_file}\n{pdf_file}\n\nTiempo: {tiempo} segundos"
        )

    except Exception as e:
        messagebox.showerror("Error", str(e))
        estado.set("Error ❌")
        progreso.set(0)
    finally:
        # Limpiar COM
        pythoncom.CoUninitialize()
        # 🔓 Desbloquear botón
        btn_generar.config(state="normal")

def ejecutar():
    threading.Thread(target=generar).start()

# ==============================
# CARGAR TXT BOTÓN
# ==============================
def cargar_archivo():
    ruta = filedialog.askopenfilename(filetypes=[("TXT", "*.txt")])
    if ruta:
        with open(ruta, "r", encoding="utf-8") as f:
            txt_input.delete("1.0", tk.END)
            txt_input.insert(tk.END, f.read())

# ==============================
# UI
# ==============================
root = TkinterDnD.Tk()
root.title("Generador Cuenta PRO")
root.geometry("600x600")

# TXT
tk.Label(root, text="Arrastra o pega el TXT").pack()

txt_input = tk.Text(root, height=10)
txt_input.pack(fill="x", padx=10)

# Activar drag & drop
txt_input.drop_target_register(DND_FILES)
txt_input.dnd_bind('<<Drop>>', drop)

tk.Button(root, text="Cargar TXT", command=cargar_archivo).pack(pady=5)

# VALOR
tk.Label(root, text="Valor Total").pack()
valor_entry = tk.Entry(root)
valor_entry.pack()

# CALENDARIO
tk.Label(root, text="Fecha").pack()
calendario = DateEntry(root, date_pattern="dd/mm/yyyy")
calendario.pack(pady=5)

# BOTON
btn_generar = tk.Button(root, text="GENERAR", command=ejecutar, bg="green", fg="white")
btn_generar.pack(pady=10)

# PROGRESO
progreso = tk.IntVar()
barra = ttk.Progressbar(root, variable=progreso, maximum=100)
barra.pack(fill="x", padx=10, pady=10)

estado = tk.StringVar()
estado.set("Arrastra un TXT o pégalo")
tk.Label(root, textvariable=estado).pack()

root.mainloop()