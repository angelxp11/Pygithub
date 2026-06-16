import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import platform
from datetime import datetime



# ---------------- UTILIDAD BASE ---------------- #

def obtener_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


# ---------------- FUNCIONES BASE ---------------- #

def leer_hojas(archivo_excel):
    wb = openpyxl.load_workbook(archivo_excel, data_only=True)
    return wb.sheetnames


def valor_celda(celda):
    valor = celda.value

    if valor is None:
        return ""

    # Si Excel la interpretó como fecha
    if isinstance(valor, datetime):
        formato = celda.number_format.lower()

        # Si el formato contiene barras y no parece fecha real,
        # probablemente era una fracción tipo 1/2
        if "/" in formato and "yy" not in formato and "dd" not in formato:
            return f"{valor.month}/{valor.day}"

        return valor.strftime("%d/%m/%Y")

    return str(valor).strip()


def leer_preguntas_desde_excel(archivo_excel, nombre_hoja):
    wb = openpyxl.load_workbook(archivo_excel, data_only=True)
    hoja = wb[nombre_hoja]

    preguntas = []

    for fila in hoja.iter_rows(min_row=2):
        if fila[0].value is None:
            continue

        enunciado = valor_celda(fila[0])
        a = valor_celda(fila[1])
        b = valor_celda(fila[2])
        c = valor_celda(fila[3])
        d = valor_celda(fila[4])
        correcta = valor_celda(fila[5]).upper()

        opciones = {
            "A": a,
            "B": b,
            "C": c,
            "D": d
        }

        if correcta not in opciones:
            raise ValueError(f"La respuesta correcta '{correcta}' no es válida")

        aiken = f"{enunciado}\n"
        for letra, texto in opciones.items():
            aiken += f"{letra}. {texto}\n"
        aiken += f"ANSWER: {correcta}\n\n"

        preguntas.append((enunciado, opciones, correcta, aiken))

    return preguntas


def generar_archivo_aiken(preguntas, ruta_salida):
    with open(ruta_salida, "w", encoding="utf-8") as f:
        for _, _, _, aiken in preguntas:
            f.write(aiken)


def generar_documento_word(preguntas, ruta_salida, xs=None):
    base_dir = obtener_base_dir()
    plantilla = os.path.join(base_dir, "plantilla.docx")

    if not os.path.exists(plantilla):
        raise FileNotFoundError("No se encontró el archivo 'plantilla.docx' junto al ejecutable")

    doc = Document(plantilla)

    if xs:
        titulo = doc.add_heading(f"EVALUACIÓN {xs}", level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Instrucciones:")
    doc.add_paragraph(
        "1. Lee cada pregunta.\n"
        "2. Marca una sola respuesta.\n"
        "3. No uses el celular.\n"
        "4. Revisa antes de entregar."
    )

    for i, (enunciado, opciones, _, _) in enumerate(preguntas, 1):
        doc.add_heading(f"Pregunta {i}", level=2)
        doc.add_paragraph(enunciado)

        for k in ["A", "B", "C", "D"]:
            doc.add_paragraph(f"{k}) {opciones[k]}")

    doc.save(ruta_salida)


def abrir_carpeta(ruta):
    if platform.system() == "Windows":
        os.startfile(ruta)
    elif platform.system() == "Darwin":
        subprocess.Popen(["open", ruta])
    else:
        subprocess.Popen(["xdg-open", ruta])


# ---------------- INTERFAZ ---------------- #

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Evaluaciones")
        self.root.geometry("500x500")
        self.root.resizable(False, False)

        self.archivo = ""

        style = ttk.Style()
        style.theme_use("clam")

        frame = ttk.Frame(root, padding=20)
        frame.pack(expand=True, fill="both")

        ttk.Label(frame, text="Generador de Evaluaciones", font=("Arial", 16, "bold")).pack(pady=10)

        ttk.Button(frame, text="Seleccionar Excel", command=self.seleccionar_archivo).pack(pady=10)

        self.label_archivo = ttk.Label(frame, text="Ningún archivo seleccionado")
        self.label_archivo.pack()

        ttk.Label(frame, text="Seleccionar hoja:").pack(pady=10)
        self.combo_hojas = ttk.Combobox(frame, state="readonly")
        self.combo_hojas.pack()

        ttk.Label(frame, text="Nombre de la evaluación:").pack(pady=10)
        self.entry_xs = ttk.Entry(frame)
        self.entry_xs.pack()

        ttk.Button(frame, text="Generar archivos", command=self.generar).pack(pady=20)

    def seleccionar_archivo(self):
        archivo = filedialog.askopenfilename(
            filetypes=[("Archivos de Excel", "*.xlsx *.xlsm")]
        )

        if archivo:
            self.archivo = archivo
            self.label_archivo.config(text=os.path.basename(archivo))

            hojas = leer_hojas(archivo)
            self.combo_hojas["values"] = hojas

            if hojas:
                self.combo_hojas.current(0)

    def generar(self):
        try:
            if not self.archivo:
                messagebox.showerror("Error", "Debes seleccionar un archivo de Excel")
                return

            hoja = self.combo_hojas.get()
            xs = self.entry_xs.get()

            if not xs:
                messagebox.showerror("Error", "Debes ingresar el nombre de la evaluación")
                return

            preguntas = leer_preguntas_desde_excel(self.archivo, hoja)

            base_dir = obtener_base_dir()

            bd_dir = os.path.join(base_dir, "bd")
            lapiz_dir = os.path.join(base_dir, "lapiz")

            os.makedirs(bd_dir, exist_ok=True)
            os.makedirs(lapiz_dir, exist_ok=True)

            nombre_base = f"EVALUACIÓN {xs}"
            archivo_aiken = os.path.join(bd_dir, nombre_base + ".txt")
            archivo_docx = os.path.join(lapiz_dir, nombre_base + ".docx")

            generar_archivo_aiken(preguntas, archivo_aiken)
            generar_documento_word(preguntas, archivo_docx, xs)

            messagebox.showinfo("Éxito", "Archivos creados correctamente")

            abrir_carpeta(base_dir)

        except Exception as e:
            messagebox.showerror("Error", str(e))


# ---------------- MAIN ---------------- #

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()